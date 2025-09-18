#!/usr/bin/env python
"""
Process Word Documents for PII Detection and Replacement

A streamlined tool for detecting and replacing PII, PHI, and PCI in Word documents.

Example Usage:
    # Detect all PII types (default)
    python process_docx.py detect sample.docx

    # Detect specific PII types (use underscores for multi-word types)
    python process_docx.py detect sample.docx --pii-types name ssn phone_number email_address

    # Replace PII with fake data
    python process_docx.py replace sample.docx --output anonymized.docx

    # Use predefined groups
    python process_docx.py detect sample.docx --pii  # Personal identifiers
    python process_docx.py detect sample.docx --phi  # Health information
    python process_docx.py detect sample.docx --pci  # Payment card data
"""

import argparse
import json
import sys
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from gliner import GLiNER

try:
    from faker import Faker
    FAKER_AVAILABLE = True
    fake = Faker()
except ImportError:
    FAKER_AVAILABLE = False
    fake = None


# All available PII types
ALL_PII_TYPES = [
    'number', 'location address street', 'pin', 'name medical professional', 'accounts',
    'policy number', 'marital status', 'passport number', 'ssn', 'discharge date', 'occupation',
    'date', 'origin', 'test result', 'name', 'location zip', 'gender', 'organization medical facility',
    'esidno', 'zip', 'date interval', 'dob', 'rate', 'organization', 'location state',
    'confirmation number', 'name given', 'time', 'cvv', 'month', 'credit card', 'planduration',
    'filename', 'age', 'numerical pii', 'money', 'physical attribute', 'address',
    'credit card expiration', 'account number', 'location', 'language',
    'location city', 'duration', 'password', 'medical process', 'county', 'phone number',
    'condition', 'email address', 'location address', 'name family', 'location country'
]

# PII group (personally identifiable information)
PII_GROUP = [
    'name', 'name given', 'name family', 'phone number', 'email address',
    'ssn', 'dob', 'age', 'gender', 'marital status', 'origin',
    'location address', 'location address street', 'location city',
    'location state', 'location country', 'location zip', 'location',
    'address', 'zip', 'county', 'passport number', 'occupation',
    'language', 'physical attribute', 'password', 'filename',
    'date', 'time', 'duration', 'date interval', 'month', 'number',
    'numerical pii', 'esidno', 'confirmation number'
]

# PHI group (protected health information)
PHI_GROUP = [
    'name', 'name medical professional', 'dob', 'age', 'gender',
    'phone number', 'email address', 'ssn', 'location address',
    'organization medical facility', 'condition', 'medical process',
    'test result', 'discharge date', 'policy number', 'account number',
    'location city', 'location state', 'location zip'
]

# PCI group (payment card industry)
PCI_GROUP = [
    'credit card', 'credit card expiration', 'cvv', 'account number',
    'accounts', 'pin', 'money', 'rate', 'planduration'
]

# PII to Faker mapping
PII_TO_FAKER = {
    "account number": lambda: fake.ean() if fake else "ACC1234567890",
    "accounts": lambda: fake.ean() if fake else "ACC1234567890",
    "address": lambda: fake.address().replace('\n', ', ') if fake else "123 Main St, City, State 12345",
    "age": lambda: str(fake.random_int(min=18, max=90)) if fake else "35",
    "condition": lambda: fake.random_element(['Hypertension', 'Diabetes', 'Asthma']) if fake else "Medical Condition",
    "confirmation number": lambda: fake.ean8() if fake else "CONF12345",
    "county": lambda: f"{fake.city()} County" if fake else "Sample County",
    "credit card": lambda: fake.credit_card_number() if fake else "4111-1111-1111-1111",
    "credit card expiration": lambda: fake.credit_card_expire() if fake else "12/25",
    "cvv": lambda: fake.credit_card_security_code() if fake else "123",
    "date": lambda: fake.date() if fake else "2024-01-01",
    "date interval": lambda: f"{fake.date()} to {fake.date()}" if fake and fake else "2024-01-01 to 2024-12-31",
    "discharge date": lambda: fake.date() if fake else "2024-03-15",
    "dob": lambda: fake.date_of_birth().strftime('%m/%d/%Y') if fake else "01/01/1990",
    "duration": lambda: f"{fake.random_int(1, 100)} {fake.random_element(['days', 'weeks', 'months'])}" if fake else "30 days",
    "email address": lambda: fake.email() if fake else "user@example.com",
    "esidno": lambda: fake.ean13() if fake else "ESI123456789",
    "filename": lambda: fake.file_name() if fake else "document.pdf",
    "gender": lambda: fake.random_element(['Male', 'Female', 'Non-binary']) if fake else "Person",
    "language": lambda: fake.language_name() if fake else "English",
    "location": lambda: fake.city() if fake else "City Name",
    "location address": lambda: fake.address().replace('\n', ', ') if fake else "123 Main St, City, State 12345",
    "location address street": lambda: fake.street_address() if fake else "123 Main Street",
    "location city": lambda: fake.city() if fake else "Anytown",
    "location country": lambda: fake.country() if fake else "United States",
    "location state": lambda: fake.state() if fake else "California",
    "location zip": lambda: fake.postcode() if fake else "12345",
    "marital status": lambda: fake.random_element(['Single', 'Married', 'Divorced']) if fake else "Status",
    "medical process": lambda: fake.random_element(['Surgery', 'X-Ray', 'MRI', 'Blood Test']) if fake else "Medical Procedure",
    "money": lambda: f"${fake.random_int(100, 100000):,}" if fake else "$50,000",
    "month": lambda: fake.month_name() if fake else "January",
    "name": lambda: fake.name() if fake else "John Doe",
    "name family": lambda: fake.last_name() if fake else "Doe",
    "name given": lambda: fake.first_name() if fake else "John",
    "name medical professional": lambda: f"Dr. {fake.name()}" if fake else "Dr. Smith",
    "number": lambda: str(fake.random_int(1, 999999)) if fake else "12345",
    "numerical pii": lambda: fake.ean() if fake else "NUM123456",
    "occupation": lambda: fake.job() if fake else "Professional",
    "organization": lambda: fake.company() if fake else "Example Corp",
    "organization medical facility": lambda: f"{fake.company()} Hospital" if fake else "General Hospital",
    "origin": lambda: fake.country() if fake else "Country",
    "passport number": lambda: fake.passport_number() if fake else "P12345678",
    "password": lambda: fake.password() if fake else "SecurePass123",
    "phone number": lambda: fake.phone_number() if fake else "(555) 123-4567",
    "physical attribute": lambda: fake.random_element(['Tall', 'Short', 'Athletic']) if fake else "Description",
    "pin": lambda: str(fake.random_int(1000, 9999)) if fake else "1234",
    "planduration": lambda: f"{fake.random_int(1, 36)} months" if fake else "12 months",
    "policy number": lambda: f"POL-{fake.ean8()}" if fake else "POL-12345678",
    "rate": lambda: f"{fake.random_int(1, 100)}%" if fake else "5%",
    "ssn": lambda: fake.ssn() if fake else "123-45-6789",
    "test result": lambda: fake.random_element(['Positive', 'Negative', 'Normal']) if fake else "Result",
    "time": lambda: fake.time() if fake else "12:00:00",
    "zip": lambda: fake.postcode() if fake else "12345"
}


def extract_all_text(document: Document) -> List[Dict[str, str]]:
    """Extract all text from document including tables."""
    extracted = []

    # Extract paragraphs
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip():
            extracted.append({
                'type': 'paragraph',
                'index': i,
                'text': paragraph.text
            })

    # Extract tables
    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    extracted.append({
                        'type': 'table_cell',
                        'table': table_idx,
                        'row': row_idx,
                        'column': col_idx,
                        'text': cell.text
                    })

    return extracted


def get_replacement(pii_type: str, original_text: str) -> str:
    """Get a replacement value for a PII type."""
    if pii_type in PII_TO_FAKER:
        return PII_TO_FAKER[pii_type]()
    return f"[{pii_type.upper().replace(' ', '_')}]"


def replace_pii_in_text(text: str, entities: List[Dict]) -> str:
    """Replace PII in text with fake data or redactions."""
    if not entities:
        return text

    sorted_entities = sorted(entities, key=lambda x: x.get('start', 0), reverse=True)
    modified_text = text

    for entity in sorted_entities:
        replacement = get_replacement(entity['label'], entity['text'])

        # Preserve capitalization
        if entity['text'].isupper():
            replacement = replacement.upper()
        elif entity['text'] and entity['text'][0].isupper():
            replacement = replacement[0].upper() + replacement[1:] if len(replacement) > 1 else replacement.upper()

        if 'start' in entity and 'end' in entity:
            modified_text = modified_text[:entity['start']] + replacement + modified_text[entity['end']:]
        else:
            modified_text = modified_text.replace(entity['text'], replacement, 1)

    return modified_text


def normalize_pii_types(pii_types: List[str]) -> List[str]:
    """Convert user-friendly underscore format to internal space format."""
    return [pii_type.replace('_', ' ') for pii_type in pii_types]


def get_pii_types(args) -> List[str]:
    """Determine which PII types to use based on arguments."""
    if args.pii_types:
        # User specified exact types - convert underscores to spaces
        return normalize_pii_types(args.pii_types)
    elif args.pii:
        return PII_GROUP
    elif args.phi:
        return PHI_GROUP
    elif args.pci:
        return PCI_GROUP
    else:
        # Default to all types
        return ALL_PII_TYPES


def cmd_detect(args):
    """Detect PII in document."""
    if not Path(args.input).exists():
        print(f"Error: Document '{args.input}' not found")
        return 1

    pii_types = get_pii_types(args)

    print(f"Loading model: {args.model}...")
    model = GLiNER.from_pretrained(args.model)

    print(f"Processing: {args.input}")
    print(f"Detecting {len(pii_types)} PII type(s)")
    if args.pii:
        print("Using PII group")
    elif args.phi:
        print("Using PHI group")
    elif args.pci:
        print("Using PCI group")
    elif not args.pii_types:
        print("Using all PII types")

    document = Document(args.input)
    content = extract_all_text(document)

    results = {
        'document': args.input,
        'total_elements': len(content),
        'pii_found': [],
        'summary': {
            'total_pii_instances': 0,
            'pii_by_type': {}
        }
    }

    # Process each element
    for element in content:
        entities = model.predict_entities(element['text'], pii_types, threshold=args.threshold)

        if entities:
            element_result = {
                'element_type': element['type'],
                'text': element['text'],
                'entities': entities
            }

            if args.format == 'redacted':
                element_result['redacted'] = replace_pii_in_text(element['text'], entities)

            results['pii_found'].append(element_result)

            for entity in entities:
                pii_type = entity['label']
                results['summary']['total_pii_instances'] += 1
                if pii_type not in results['summary']['pii_by_type']:
                    results['summary']['pii_by_type'][pii_type] = []
                results['summary']['pii_by_type'][pii_type].append(entity['text'])

    # Output results
    if args.format == 'summary':
        print("\n" + "=" * 60)
        print("PII DETECTION SUMMARY")
        print("=" * 60)
        print(f"Total PII instances found: {results['summary']['total_pii_instances']}")

        if results['summary']['pii_by_type']:
            print("\nPII by type:")
            for pii_type, instances in results['summary']['pii_by_type'].items():
                unique = list(set(instances))[:5]
                print(f"\n{pii_type}: {len(instances)} instance(s)")
                for inst in unique:
                    print(f"  - {inst}")
                if len(set(instances)) > 5:
                    print(f"  ... and {len(set(instances)) - 5} more unique value(s)")
        else:
            print("\nNo PII detected with current settings.")

    elif args.format == 'json':
        print(json.dumps(results, indent=2))

    elif args.format == 'redacted':
        print("\n" + "=" * 60)
        print("REDACTED CONTENT")
        print("=" * 60)
        for element in results['pii_found']:
            print(f"\n[{element['element_type'].upper()}]")
            print(element['redacted'])

    # Save if requested
    if args.output:
        with open(args.output, 'w') as f:
            json.dump(results, f, indent=2)
        print(f"\nResults saved to: {args.output}")

    return 0


def cmd_replace(args):
    """Replace PII in document with fake data."""
    if not Path(args.input).exists():
        print(f"Error: Document '{args.input}' not found")
        return 1

    if not FAKER_AVAILABLE:
        print("Note: Faker library not installed. Using generic replacements.")
        print("Install with: uv pip install faker\n")

    pii_types = get_pii_types(args)

    print(f"Loading model: {args.model}...")
    model = GLiNER.from_pretrained(args.model)

    print(f"Processing: {args.input}")
    print(f"Replacing {len(pii_types)} PII type(s)")

    document = Document(args.input)
    replacements_made = 0

    # Process paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            entities = model.predict_entities(paragraph.text, pii_types, threshold=args.threshold)
            if entities:
                new_text = replace_pii_in_text(paragraph.text, entities)
                if new_text != paragraph.text:
                    paragraph.text = new_text
                    replacements_made += len(entities)

    # Process tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    entities = model.predict_entities(cell.text, pii_types, threshold=args.threshold)
                    if entities:
                        new_text = replace_pii_in_text(cell.text, entities)
                        if new_text != cell.text:
                            cell.text = new_text
                            replacements_made += len(entities)

    # Save modified document
    output_path = args.output or args.input.replace('.docx', '_anonymized.docx')
    document.save(output_path)

    print(f"\nCompleted! Made {replacements_made} replacements")
    print(f"Anonymized document saved to: {output_path}")

    if FAKER_AVAILABLE:
        print("Using Faker library for realistic replacements")
    else:
        print("Using generic replacements")

    return 0


def main():
    parser = argparse.ArgumentParser(
        prog='process_docx',
        description='Process Word documents for PII/PHI/PCI detection and replacement'
    )
    subparsers = parser.add_subparsers(dest='command', help='Available commands')

    # Detect command
    detect_parser = subparsers.add_parser('detect', help='Detect PII in document')
    detect_parser.add_argument('input', help='Input Word document (.docx)')
    detect_parser.add_argument('--model', default='wordcab/wordcab-pii-detection-large-v0.2',
                              help='Model name or path')
    detect_parser.add_argument('--threshold', type=float, default=0.5,
                              help='Detection threshold (default: 0.5)')

    # PII type selection (mutually exclusive groups)
    pii_group = detect_parser.add_mutually_exclusive_group()
    pii_group.add_argument('--pii-types', nargs='+',
                          help='Specific PII types to detect (e.g., name ssn phone_number credit_card)')
    pii_group.add_argument('--pii', action='store_true',
                          help='Detect PII group (personal identifiers)')
    pii_group.add_argument('--phi', action='store_true',
                          help='Detect PHI group (health information)')
    pii_group.add_argument('--pci', action='store_true',
                          help='Detect PCI group (payment card data)')
    detect_parser.add_argument('--format', choices=['summary', 'json', 'redacted'],
                              default='summary', help='Output format')
    detect_parser.add_argument('--output', help='Save results to file')

    # Replace command
    replace_parser = subparsers.add_parser('replace', help='Replace PII with fake data')
    replace_parser.add_argument('input', help='Input Word document (.docx)')
    replace_parser.add_argument('--output', help='Output file path')
    replace_parser.add_argument('--model', default='knowledgator/gliner-multitask-large-v0.5',
                               help='Model name or path')
    replace_parser.add_argument('--threshold', type=float, default=0.3,
                               help='Detection threshold')

    # PII type selection (mutually exclusive groups)
    replace_pii_group = replace_parser.add_mutually_exclusive_group()
    replace_pii_group.add_argument('--pii-types', nargs='+',
                                  help='Specific PII types to replace (e.g., name ssn phone_number credit_card)')
    replace_pii_group.add_argument('--pii', action='store_true',
                                  help='Replace PII group')
    replace_pii_group.add_argument('--phi', action='store_true',
                                  help='Replace PHI group')
    replace_pii_group.add_argument('--pci', action='store_true',
                                  help='Replace PCI group')

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return 1

    # Execute command
    if args.command == 'detect':
        return cmd_detect(args)
    elif args.command == 'replace':
        return cmd_replace(args)


if __name__ == "__main__":
    sys.exit(main())